from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\yulan\source2\repos\classlift\classlift_management\Docs\ClassLift.com Cloudflare Setup and Redirect Guide.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

PURPLE = "5146E5"
NAVY = "14213D"
BLUE = "2E74B5"
MUTED = "667085"
LIGHT = "F3F4FF"
PALE = "F7F8FA"
GREEN = "137A55"
AMBER = "8A5A00"
RED = "A12828"
WHITE = "FFFFFF"
BORDER = "D9DCE6"

doc = Document()
sec = doc.sections[0]
sec.different_first_page_header_footer = False
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in [
    ("Title", 28, NAVY, 0, 8),
    ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 17, PURPLE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 5),
    ("Heading 3", 11.5, NAVY, 9, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_style in ["List Bullet", "List Number"]:
    st = styles[list_style]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.18

if "Code Block" not in [s.name for s in styles]:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["Code Block"]
code_style.font.name = "Consolas"
code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor.from_string(NAVY)
code_style.paragraph_format.left_indent = Inches(0.18)
code_style.paragraph_format.right_indent = Inches(0.18)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(7)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[i] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def add_run(p, text, bold=False, color=None, size=None, font=None):
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    if font:
        r.font.name = font
        r._element.rPr.rFonts.set(qn("w:ascii"), font)
        r._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    return r

def callout(label, text, fill=LIGHT, color=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, label + "  ", bold=True, color=color)
    add_run(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def code(text):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE)
    pPr.append(shd)
    add_run(p, text, font="Consolas", size=9)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = False
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])

# Running furniture (populate both default and even-page parts for renderer compatibility)
for header in [sec.header, sec.even_page_header]:
    hp = header.paragraphs[0]
    hp.text = "CLASSLIFT  |  DOMAIN SETUP GUIDE"
    hp.style = normal
    hp.paragraph_format.space_after = Pt(0)
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.bold = True
    hp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

for footer in [sec.footer, sec.even_page_footer]:
    fp = footer.paragraphs[0]
    add_run(fp, "ClassLift Cloudflare Guide   •   Page ", color=MUTED, size=8.5)
    add_page_number(fp)
    for r in fp.runs:
        r.font.color.rgb = RGBColor.from_string(MUTED)
        r.font.size = Pt(8.5)

# Cover / opening block
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after = Pt(4)
add_run(p, "OPERATIONS GUIDE", bold=True, color=PURPLE, size=10)

p = doc.add_paragraph(style="Title")
p.add_run("Set Up classlift.com and Redirect It to classlift.ca")
p = doc.add_paragraph(style="Subtitle")
p.add_run("Cloudflare Free Plan • Permanent 301 redirect • Path and query preservation")

callout(
    "Recommended outcome",
    "Keep classlift.ca as the primary Canadian website. Route classlift.com and www.classlift.com through Cloudflare and permanently redirect visitors to the matching page on classlift.ca.",
)

doc.add_heading("What this guide will configure", level=1)
for t in [
    "Add classlift.com to a free Cloudflare account.",
    "Move authoritative DNS to Cloudflare without losing existing email records.",
    "Create proxied DNS records for classlift.com and www.classlift.com.",
    "Create one 301 redirect rule that preserves page paths and query strings.",
    "Verify HTTPS, redirect behavior, and common failure conditions.",
]:
    bullet(t)

doc.add_heading("Final traffic flow", level=2)
code("classlift.com/*  →  Cloudflare edge  →  301  →  classlift.ca/same-path")

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [4680, 4680])
headers = ["Visitor enters", "Visitor reaches"]
for i, h in enumerate(headers):
    shade(table.rows[0].cells[i], PURPLE)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, h, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for left, right in [
    ("https://classlift.com/", "https://classlift.ca/"),
    ("https://www.classlift.com/pricing.html", "https://classlift.ca/pricing.html"),
    ("https://classlift.com/contact.html?source=email", "https://classlift.ca/contact.html?source=email"),
]:
    row = table.add_row()
    row.cells[0].text = left
    row.cells[1].text = right
set_table_geometry(table, [4680, 4680])

doc.add_page_break()
doc.add_heading("Before you begin", level=1)
callout(
    "Important",
    "Changing nameservers affects every DNS service for classlift.com, including email. Take screenshots or export the current DNS zone before making changes.",
    fill="FFF5E5",
    color=AMBER,
)

doc.add_heading("You will need", level=2)
for t in [
    "Login access to the registrar where classlift.com was purchased.",
    "Access to the email address used for the Cloudflare account.",
    "A copy of all current DNS records, especially MX and TXT records.",
    "Permission to change the domain's authoritative nameservers.",
]:
    bullet(t)

doc.add_heading("DNS records that must be protected", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [1400, 2700, 5260])
for i, h in enumerate(["Type", "Usually used for", "What to do"]):
    shade(table.rows[0].cells[i], PURPLE)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, h, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for values in [
    ("MX", "Receiving email", "Copy exactly; do not delete unless confirmed obsolete."),
    ("TXT", "SPF, DKIM, DMARC, verification", "Copy the complete value exactly, including punctuation."),
    ("CNAME", "Websites and service verification", "Identify whether it is still needed before replacing it."),
    ("A / AAAA", "Website hosting", "Check whether it conflicts with the new redirect-only records."),
]:
    row = table.add_row()
    for i, value in enumerate(values):
        row.cells[i].text = value
set_table_geometry(table, [1400, 2700, 5260])

doc.add_heading("Step 1 — Create and secure the Cloudflare account", level=1)
for t in [
    "Open https://dash.cloudflare.com/ and create an account using an email address you control.",
    "Verify the account email address.",
    "Enable two-factor authentication under the account profile or authentication settings.",
]:
    bullet(t)

doc.add_heading("Step 2 — Add classlift.com", level=1)
for t in [
    "In the Cloudflare dashboard, choose Add a domain, Add site, or Onboard a domain. Cloudflare may use slightly different wording.",
    "Enter classlift.com only. Do not enter https:// and do not enter www.",
    "Choose the Free plan and continue.",
    "Allow Cloudflare to scan the current DNS records.",
    "Compare the imported records with the copy made before starting. Restore any missing email records before continuing.",
]:
    bullet(t)

doc.add_page_break()
doc.add_heading("Step 3 — Change nameservers at the registrar", level=1)
p = doc.add_paragraph()
p.add_run("Cloudflare will assign two nameservers unique to classlift.com, similar to:")
code("example-one.ns.cloudflare.com\nexample-two.ns.cloudflare.com")
callout("Do not copy the examples", "Use only the two nameservers displayed inside your own Cloudflare dashboard.", fill="FFF5E5", color=AMBER)

for t in [
    "Log in to the registrar where classlift.com was purchased.",
    "Open Domain Settings, Nameservers, or DNS Management.",
    "Choose the option to use custom nameservers.",
    "Replace the existing nameservers with the two nameservers assigned by Cloudflare.",
    "Save the change and return to Cloudflare.",
    "Select Check nameservers now and wait until the zone status becomes Active.",
]:
    bullet(t)

callout(
    "DNSSEC note",
    "If the registrar reports a DNSSEC conflict, follow Cloudflare's instruction to disable the old DNSSEC configuration before changing nameservers. Re-enable DNSSEC from Cloudflare after activation if desired.",
    fill=PALE,
    color=BLUE,
)

doc.add_heading("Step 4 — Create redirect-only DNS records", level=1)
p = doc.add_paragraph()
p.add_run("After classlift.com shows ")
add_run(p, "Active", bold=True, color=GREEN)
p.add_run(", open DNS → Records in the classlift.com zone.")

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [1050, 1000, 3400, 1850, 2060])
for i, h in enumerate(["Type", "Name", "Target", "Proxy", "TTL"]):
    shade(table.rows[0].cells[i], PURPLE)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, h, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for values in [
    ("A", "@", "192.0.2.1", "Proxied", "Auto"),
    ("CNAME", "www", "classlift.com", "Proxied", "Auto"),
]:
    row = table.add_row()
    for i, value in enumerate(values):
        row.cells[i].text = value
set_table_geometry(table, [1050, 1000, 3400, 1850, 2060])

for t in [
    "The orange cloud must be enabled for both records; this is the Proxied state.",
    "192.0.2.1 is a reserved documentation address. Cloudflare should execute the redirect before any request reaches that address.",
    "Remove or replace conflicting @ or www web records only after confirming they are not required.",
    "Do not remove MX, SPF, DKIM, DMARC, or unrelated verification records.",
]:
    bullet(t)

step5_heading = doc.add_heading("Step 5 — Create the 301 redirect rule", level=1)
step5_heading.paragraph_format.page_break_before = True
p = doc.add_paragraph()
p.add_run("Inside the ")
add_run(p, "classlift.com", bold=True)
p.add_run(" zone, open one of these paths depending on the current Cloudflare interface:")
code("Rules → Redirect Rules → Create rule\nRules → Overview → Redirect Rules → Single Redirects → Create rule")

doc.add_heading("Rule configuration", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [2700, 6660])
for i, h in enumerate(["Field", "Value"]):
    shade(table.rows[0].cells[i], PURPLE)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, h, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
rows = [
    ("Rule name", "Redirect classlift.com to classlift.ca"),
    ("Match type", "Custom filter expression"),
    ("Incoming request expression", '(http.host eq "classlift.com") or (http.host eq "www.classlift.com")'),
    ("Redirect type", "Dynamic"),
    ("Destination expression", 'concat("https://classlift.ca", http.request.uri.path)'),
    ("Status code", "301 — Permanent Redirect"),
    ("Query string", "Preserve query string: enabled"),
]
for label, value in rows:
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
set_table_geometry(table, [2700, 6660])

doc.add_heading("What the destination expression means", level=2)
code('concat("https://classlift.ca", http.request.uri.path)')
p = doc.add_paragraph()
p.add_run("The expression joins the destination host ")
add_run(p, "https://classlift.ca", bold=True)
p.add_run(" with the visitor's original path. It prevents every request from being sent only to the home page.")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [3800, 2200, 3360])
for i, h in enumerate(["Original request", "Original path", "Destination"]):
    shade(table.rows[0].cells[i], BLUE)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, h, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for values in [
    ("classlift.com/", "/", "classlift.ca/"),
    ("classlift.com/pricing.html", "/pricing.html", "classlift.ca/pricing.html"),
    ("classlift.com/contact.html", "/contact.html", "classlift.ca/contact.html"),
]:
    row = table.add_row()
    for i, value in enumerate(values):
        row.cells[i].text = value
set_table_geometry(table, [3800, 2200, 3360])

callout(
    "About #pricing",
    "The part after # is a browser fragment and is not sent to Cloudflare. When the redirect destination does not supply a different fragment, the browser normally retains it.",
    fill=PALE,
    color=BLUE,
)

doc.add_page_break()
doc.add_heading("Step 6 — Verify HTTPS and redirect behavior", level=1)
for t in [
    "Wait for Cloudflare to issue edge certificates for classlift.com and www.classlift.com.",
    "Open a private or incognito browser window to avoid cached redirect results.",
    "Test HTTP and HTTPS for both the root and www hostnames.",
    "Test a real path, a query string, and the pricing-page fragment.",
]:
    bullet(t)

doc.add_heading("Browser test checklist", level=2)
for t in [
    "http://classlift.com redirects to https://classlift.ca/",
    "https://classlift.com redirects without a certificate warning",
    "https://www.classlift.com redirects correctly",
    "https://classlift.com/pricing.html redirects to the matching .ca page",
    "https://classlift.com/pricing.html?source=test retains ?source=test",
    "https://classlift.com/pricing.html#pricing opens the pricing section on classlift.ca",
]:
    bullet(t)

doc.add_heading("PowerShell verification", level=2)
code("curl.exe -I https://classlift.com/\ncurl.exe -I https://www.classlift.com/pricing.html\ncurl.exe -I \"https://classlift.com/contact.html?source=test\"")
p = doc.add_paragraph()
p.add_run("A correct response should include:")
code("HTTP/2 301\nlocation: https://classlift.ca/pricing.html")

doc.add_heading("Troubleshooting", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [3250, 6110])
for i, h in enumerate(["Problem", "Checks and corrective action"]):
    shade(table.rows[0].cells[i], PURPLE)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, h, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
problems = [
    ("Redirect Rules is missing", "Confirm that you selected the classlift.com website zone, not the account home page. The zone should show Active."),
    ("Cloudflare says nameservers are pending", "Confirm that both assigned Cloudflare nameservers replaced the old nameservers at the registrar. Allow time for propagation."),
    ("Redirect rule does not run", "Confirm that @ and www DNS records exist and show the orange Proxied cloud. Review the hostname expression."),
    ("Every request goes to the home page", "Use the dynamic concat expression instead of a fixed destination URL."),
    ("Query string disappears", "Enable Preserve query string in the redirect rule."),
    ("Email stops working", "Restore the original MX and email-related TXT records. DNS web redirects should not require changing those records."),
    ("Browser shows the old result", "Test in a private window, clear DNS/browser cache, or use curl.exe -I to inspect the current response."),
    ("HTTPS certificate warning", "Wait for Cloudflare edge-certificate issuance and confirm both hostnames are proxied."),
]
for problem, action in problems:
    row = table.add_row()
    row.cells[0].text = problem
    row.cells[1].text = action
set_table_geometry(table, [3250, 6110])

doc.add_heading("Completion checklist", level=1)
for t in [
    "Cloudflare account has two-factor authentication enabled.",
    "classlift.com shows Active in Cloudflare.",
    "Existing MX and email TXT records remain intact.",
    "@ and www DNS records are Proxied.",
    "The 301 redirect rule is deployed.",
    "Paths and query strings are preserved.",
    "HTTP, HTTPS, root, www, and pricing-page tests pass.",
    "classlift.ca remains the canonical primary website.",
]:
    bullet(t)

callout(
    "Operational recommendation",
    "Keep classlift.com registered and renewed even while classlift.ca remains primary. Review the redirect after any DNS-provider or website-hosting change.",
    fill=LIGHT,
    color=PURPLE,
)

# Document properties
doc.core_properties.title = "Set Up classlift.com and Redirect It to classlift.ca"
doc.core_properties.subject = "Cloudflare domain onboarding and redirect guide"
doc.core_properties.author = "ClassLift"
doc.core_properties.keywords = "ClassLift, Cloudflare, DNS, redirect, classlift.com, classlift.ca"

# Use the same running header/footer on odd and even pages.
settings = doc.settings._element
even_odd = settings.find(qn("w:evenAndOddHeaders"))
if even_odd is not None:
    settings.remove(even_odd)

doc.save(OUT)
print(str(OUT))
